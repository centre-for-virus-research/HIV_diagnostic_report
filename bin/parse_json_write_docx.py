#!/usr/bin/env python3.6


import json
import argparse
import os as os
from docx import Document
import datetime 

if __name__ == "__main__":

    sample2subtype = {}
    parser = argparse.ArgumentParser()
    parser.add_argument('--json', required=True, help='input json file containing query results')
    parser.add_argument('--output', required=False, help='name of tab-delimited text file containing sample subtypes')
    parser.add_argument('--reports', required=False, action='store_true', help='if this flag is included, .docx reports will be produced for each sample')
    args = parser.parse_args()

    json_in = ""
    subtype_output = ""

    if args.json:
        json_in = args.json
    else:
        print ("json file cannot be read by parser")

    if args.output:
        subtype_output = args.output
    else:
        subtype_output = "subtypes.txt"

    now = datetime.datetime.now()
    if now.month < 10:
        time_now = (str(now.day) + "-0" +str(now.month) + "-" + str(now.year))
    else:
        time_now =(str(now.day) + "-" +str(now.month) + "-" + str(now.year))

    path = time_now + "_reports/"
    if args.reports:
        try:
            os.makedirs(os.path.dirname(path))
        except OSError as ecx:
            if exc.errno != errno.EEXIST:
                raise

# parse the sierrapy json output for relevant information - subtype and DRMs
    with open(json_in) as json_file:
        data = json.load(json_file)

        for i in data:
            sample = i['inputSequence']['header']
            report_file_name = path +  sample + "_report.docx"
            document = Document()
            document.add_heading("Sequence Summary", level=1)
# gene name and codon information
            for j in i["alignedGeneSequences"]:
                start = j["firstAA"]
                end = j["lastAA"]
                gene = j["gene"]["name"]

                p1 = document.add_paragraph("Sequence includes " + gene +": codons " + str(start) + "-" + str(end))
# subtype information
            subtype = i['subtypeText']
            subtype_id = str.split(subtype, "(")[0]
            subtype_message = ""
            if subtype == 'NA':
                subtype_message = "No subtype information for sample:\t" + sample
            elif subtype_id.rstrip() in("B", "C"):
                sample2subtype[sample] = subtype
                subtype_message = "Subtype: "+subtype
            else:
                #print ("Complex subtype. Sample will be passed to SCUEAL:\t" + sample + "\t" + subtype)
                subtype_message = subtype + "\tThis info will be updated by SCUEAL"
                sample2subtype[sample] = subtype
            p2 = document.add_paragraph(subtype_message)

# Drug resistance information
            for j in i["drugResistance"]:
                currentGene = j["gene"]["name"]
                document.add_heading ("Drug Resistance Interpretation: " + currentGene + "\n", level=1)
                mutations_dict = {}
                for k in j["mutationsByTypes"]:
                    mutation_type = k["mutationType"]
                    mutation_list = []
                    for m in k["mutations"]:
                        for key in m:
                            mutation_list.append(m[key])
                    mutation_string = ""
                    if not mutation_list:
                        mutation_string = "None"
                    else:
                        mutation_string = ", ".join(mutation_list)
                    mutations_dict[mutation_type] = mutation_string


                scores_dict = {}
                numerical_scores_dict = {}
                drug_class = ""
                drug_name = ""
                drug_abbr = ""
                drug_fullname = ""
                drug_score = ""
                drug_text = ""
                drug_id = ""
                drug_value = ""
                m_info = {}


                if currentGene == 'RT':
                    N = {}
                    NN = {}
                    for key, value in mutations_dict.items():
                        if key == 'Other':
                            p3 = document.add_paragraph(key + " Mutations: " + value)
                        else:
                            p3 = document.add_paragraph(key + " Resistance Mutations: " + value)
                    for k in j["drugScores"]:
                        drug_class = k["drugClass"]["name"]
                        drug_name = k["drug"]["name"]
                        drug_abbr = k["drug"]["displayAbbr"]
                        drug_fullname = k["drug"]["fullName"]
                        drug_score = k["score"]
                        drug_text = k["text"]

                        drug_id = drug_fullname + " (" + drug_abbr + ") "
                        drug_value = [drug_score, drug_text]
                        if drug_class == "NRTI":
                            N[drug_id] = drug_value
                        elif drug_class =="NNRTI":
                            NN[drug_id] = drug_value

                        m_name = ""
                        m_type = ""
                        m_text = ""
                        if drug_score != 0.0:

                            for p in k["partialScores"]:
                                for key, value in p.items():
                                    if key =='mutations':
                                        for x, y in value[0].items():
                                            if x == 'text':
                                                m_name = y
                                            if x == 'comments':
                                                m_text = y[0]["text"]
                                            m_info[m_name] = m_text
                    document.add_heading("Nucleoside Reverse Transcriptase Inhibitors", level=2)
                    p4 = document.add_paragraph()
                    for key, value in N.items():
                        p4.add_run(key).bold = True
                        p4.add_run("\t" + value[1] + "\n")

                    document.add_heading("Non-Nucleoside Reverse Transcriptase Inhibitors", level=2)
                    p5 = document.add_paragraph()
                    for key, value in NN.items():
                        p5.add_run(key).bold = True
                        p5.add_run("\t" + value[1] + "\n")

                    document.add_heading ("Mutation Scoring: " + currentGene + "\n", level=1)
                    p7 = document.add_paragraph("Nucleoside Reverse Transcriptase Inhibitors\n")
                    for key, value in N.items():
                        p7.add_run(key).bold = True
                        score  = value[0]
                        if score != 0.0:
                            p7.add_run("\t" + str(score) + "\n").bold=True
                        else:
                            p7.add_run("\t" + str(score) + "\n")

                    p8 = document.add_paragraph("Non-Nucleoside Reverse Transcriptase Inhibitors\n")
                    for key, value in NN.items():
                        p8.add_run(key).bold = True
                        score  = value[0]
                        if score != 0.0:
                            p8.add_run("\t" + str(score) + "\n").bold=True
                        else:
                            p8.add_run("\t" + str(score) + "\n")
                    if not m_info:
                        continue
                    else:
                        document.add_heading(currentGene + " Comments", level=2)
                        for key, value in m_info.items():
                            p9 = document.add_paragraph(value, style='List Bullet')


                elif currentGene =='PR' or currentGene == 'IN':
                    for key, value in mutations_dict.items():
                        if key == 'Other':
                            p3 = document.add_paragraph(drug_class + " " + key + " Mutations: " + value)
                        else:
                            p3 = document.add_paragraph(drug_class + " " + key + " Resistance Mutations: " + value)


                    for k in j["drugScores"]:
                        drug_class = k["drugClass"]["name"]
                        drug_name = k["drug"]["name"]
                        drug_abbr = k["drug"]["displayAbbr"]
                        drug_fullname = k["drug"]["fullName"]
                        drug_score = k["score"]
                        drug_text = k["text"]

                        drug_id = drug_fullname + " (" + drug_abbr + ") "
                        drug_value = [drug_score, drug_text]

                        scores_dict[drug_id] = drug_value
                        # drug scores and comments
                        m_name = ""
                        m_type = ""
                        m_text = ""
                        if drug_score != 0.0:

                            for p in k["partialScores"]:
                                for key, value in p.items():
                                    if key =='mutations':
                                        for x, y in value[0].items():
                                            if x == 'text':
                                                m_name = y
                                            if x == 'comments':
                                                m_text = y[0]["text"]
                                            m_info[m_name] = m_text


                    document.add_heading(drug_class, level=2)
                    p4 = document.add_paragraph()

                    for key, value in scores_dict.items():
                        p4.add_run(key + "\t").bold = True
                        p4.add_run(value[1] + "\n")

                    document.add_heading ("Mutation Scoring: " + currentGene + "\n", level=1)
                    p6 = document.add_paragraph()
                    for key, value in scores_dict.items():
                        p6.add_run(key).bold = True
                        score  = value[0]
                        if score != 0.0:
                            p6.add_run("\t" + str(score) + "\n").bold=True
                        else:
                            p6.add_run("\t" + str(score) + "\n")

                    if not m_info:
                        continue
                    else:
                        document.add_heading(currentGene + " Comments", level=2)
                        for key, value in m_info.items():
                            p7 = document.add_paragraph(value, style='List Bullet')
                else:
                    print("Gene is not PR or RT")
            if args.reports:
               # print ("Saving report: " + report_file_name)
                document.save(report_file_name)
    with open(subtype_output, "w+") as out:
        for i in sample2subtype:
            out.write(i + "\t" + sample2subtype[i] + "\n")


